"""
Backend функция для экспорта ПАБ в формат Word (DOCX)
Возвращает файл .docx для скачивания
"""
import json
import os
from typing import Dict, Any
import psycopg2
from io import BytesIO
import base64
import requests

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Экспорт ПАБ в Word документ
    
    Query параметры:
    - ids: строка с ID через запятую (например: "1,2,3")
    """
    
    method = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'GET, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    if method != 'GET':
        return {
            'statusCode': 405,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Only GET allowed'}),
            'isBase64Encoded': False
        }
    
    query_params = event.get('queryStringParameters', {})
    ids_str = query_params.get('ids', '')
    
    if not ids_str:
        return {
            'statusCode': 400,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Parameter ids is required'}),
            'isBase64Encoded': False
        }
    
    try:
        pab_ids = [int(id_str.strip()) for id_str in ids_str.split(',')]
    except ValueError:
        return {
            'statusCode': 400,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Invalid ids format'}),
            'isBase64Encoded': False
        }
    
    database_url = os.environ.get('DATABASE_URL')
    if not database_url:
        return {
            'statusCode': 500,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Database not configured'}),
            'isBase64Encoded': False
        }
    
    try:
        conn = psycopg2.connect(database_url)
        cursor = conn.cursor()
        
        ids_placeholder = ','.join(['%s'] * len(pab_ids))
        
        query = f"""
            SELECT 
                p.id,
                p.doc_number,
                p.doc_date,
                p.inspector_fio,
                p.inspector_position,
                p.location,
                p.checked_object,
                p.department,
                p.photo_url,
                p.status
            FROM t_p80499285_psot_realization_pro.pab_records p
            WHERE p.id IN ({ids_placeholder})
            ORDER BY p.doc_date DESC
        """
        
        cursor.execute(query, pab_ids)
        pab_records = cursor.fetchall()
        
        if not pab_records:
            cursor.close()
            conn.close()
            return {
                'statusCode': 404,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'No PAB records found'}),
                'isBase64Encoded': False
            }
        
        pabs_with_obs = []
        
        for pab_row in pab_records:
            pab_id = pab_row[0]
            
            obs_query = """
                SELECT 
                    observation_number,
                    description,
                    category,
                    conditions_actions,
                    hazard_factors,
                    measures,
                    responsible_person,
                    deadline,
                    photo_url,
                    status
                FROM t_p80499285_psot_realization_pro.pab_observations
                WHERE pab_record_id = %s
                ORDER BY observation_number
            """
            
            cursor.execute(obs_query, (pab_id,))
            observations = cursor.fetchall()
            
            pabs_with_obs.append({
                'id': pab_row[0],
                'doc_number': pab_row[1],
                'doc_date': str(pab_row[2]) if pab_row[2] else '',
                'inspector_fio': pab_row[3],
                'inspector_position': pab_row[4],
                'location': pab_row[5],
                'checked_object': pab_row[6],
                'department': pab_row[7],
                'photo_url': pab_row[8] or '',
                'status': pab_row[9],
                'observations': [{
                    'observation_number': obs[0],
                    'description': obs[1],
                    'category': obs[2],
                    'conditions_actions': obs[3],
                    'hazard_factors': obs[4],
                    'measures': obs[5],
                    'responsible_person': obs[6],
                    'deadline': str(obs[7]) if obs[7] else '',
                    'photo_url': obs[8] or '',
                    'status': obs[9]
                } for obs in observations]
            })
        
        cursor.close()
        conn.close()
        
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        def download_image(url: str) -> BytesIO:
            """Скачивает изображение по URL и возвращает BytesIO объект"""
            try:
                response = requests.get(url, timeout=10)
                if response.status_code == 200:
                    return BytesIO(response.content)
            except Exception as e:
                print(f"Error downloading image from {url}: {str(e)}")
            return None
        
        def get_status_label(status: str) -> str:
            status_map = {
                'new': 'Новый',
                'in_progress': 'В работе',
                'completed': 'Выполнен',
                'overdue': 'Просрочен'
            }
            return status_map.get(status, 'Новый')
        
        doc = Document()
        
        # Настройка полей документа
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        for pab in pabs_with_obs:
            # Заголовок - ПРОТОКОЛ АУДИТА БЕЗОПАСНОСТИ
            heading1 = doc.add_heading('ПРОТОКОЛ АУДИТА БЕЗОПАСНОСТИ', level=1)
            heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading1.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            
            # Номер ПАБ
            heading2 = doc.add_heading(pab['doc_number'], level=2)
            heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading2.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            
            doc.add_paragraph()
            
            # Информационный блок в рамке (таблица)
            info_table = doc.add_table(rows=7, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Дата:', pab['doc_date']),
                ('Проверяющий:', pab['inspector_fio']),
                ('Должность:', pab['inspector_position']),
                ('Подразделение:', pab['department']),
                ('Участок:', pab['location']),
                ('Проверяемый объект:', pab['checked_object']),
                ('Статус ПАБ:', get_status_label(pab['status']))
            ]
            
            for idx, (label, value) in enumerate(info_data):
                row = info_table.rows[idx]
                label_cell = row.cells[0]
                value_cell = row.cells[1]
                
                label_para = label_cell.paragraphs[0]
                label_para.text = label
                label_para.runs[0].font.bold = True
                label_para.runs[0].font.size = Pt(10)
                
                value_para = value_cell.paragraphs[0]
                value_para.text = str(value) if value else '—'
                value_para.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()
            
            # Фото объекта проверки (если есть)
            if pab.get('photo_url'):
                photo_table = doc.add_table(rows=2, cols=1)
                photo_table.style = 'Table Grid'
                
                label_cell = photo_table.rows[0].cells[0]
                label_para = label_cell.paragraphs[0]
                label_para.text = 'Фото объекта:'
                label_para.runs[0].font.bold = True
                label_para.runs[0].font.size = Pt(10)
                
                photo_cell = photo_table.rows[1].cells[0]
                photo_para = photo_cell.paragraphs[0]
                image_stream = download_image(pab['photo_url'])
                if image_stream:
                    photo_para.add_run().add_picture(image_stream, width=Inches(4.0))
                else:
                    photo_para.text = f"Не удалось загрузить фото: {pab['photo_url']}"
                
                doc.add_paragraph()
            
            # Заголовок НАБЛЮДЕНИЯ
            obs_heading = doc.add_heading('НАБЛЮДЕНИЯ', level=2)
            obs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in obs_heading.runs:
                run.font.size = Pt(13)
                run.font.bold = True
            
            # Наблюдения
            for obs in pab['observations']:
                # Рамка для каждого наблюдения
                obs_table = doc.add_table(rows=1, cols=1)
                obs_table.style = 'Table Grid'
                obs_cell = obs_table.rows[0].cells[0]
                
                # Заголовок наблюдения
                obs_title = obs_cell.add_paragraph()
                obs_title.add_run(f"Наблюдение №{obs['observation_number']}").bold = True
                obs_title.runs[0].font.size = Pt(11)
                obs_title_status = obs_title.add_run(f"  [{get_status_label(obs['status'])}]")
                obs_title_status.font.size = Pt(9)
                
                # Поля наблюдения
                fields = [
                    ('Описание:', obs['description']),
                    ('Категория:', obs['category']),
                    ('Вид условий и действий:', obs['conditions_actions']),
                    ('Опасные факторы:', obs['hazard_factors']),
                    ('Мероприятия:', obs['measures']),
                    ('Ответственный:', obs['responsible_person'] or '—'),
                    ('Срок выполнения:', obs['deadline'] or '—'),
                ]
                
                for field_label, field_value in fields:
                    field_para = obs_cell.add_paragraph()
                    field_para.add_run(f"{field_label} ").bold = True
                    field_para.runs[0].font.size = Pt(10)
                    
                    value_run = field_para.add_run(str(field_value))
                    value_run.font.size = Pt(10)
                
                # Фото наблюдения
                if obs.get('photo_url'):
                    photo_para = obs_cell.add_paragraph()
                    photo_para.add_run('Фотография нарушения:').bold = True
                    photo_para.runs[0].font.size = Pt(10)
                    
                    image_stream = download_image(obs['photo_url'])
                    if image_stream:
                        img_para = obs_cell.add_paragraph()
                        img_para.add_run().add_picture(image_stream, width=Inches(4.0))
                    else:
                        err_para = obs_cell.add_paragraph()
                        err_para.text = f"Не удалось загрузить фото: {obs['photo_url']}"
                        err_para.runs[0].font.size = Pt(9)
                
                doc.add_paragraph()
            
            # Блок подписей
            doc.add_paragraph()
            
            # Заголовок ПОДПИСИ
            signatures_heading = doc.add_heading('ПОДПИСИ', level=2)
            signatures_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in signatures_heading.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            
            # Таблица подписей в рамке
            sig_table = doc.add_table(rows=1, cols=2)
            sig_table.style = 'Table Grid'
            
            # Левая колонка - Проверяющий
            left_cell = sig_table.rows[0].cells[0]
            
            left_label = left_cell.paragraphs[0]
            left_label.add_run('Проверяющий:').bold = True
            left_label.runs[0].font.size = Pt(10)
            
            left_name = left_cell.add_paragraph(pab['inspector_fio'])
            left_name.runs[0].font.size = Pt(10)
            
            left_cell.add_paragraph()
            left_cell.add_paragraph()
            
            left_sign = left_cell.add_paragraph('Подпись ______________')
            left_sign.runs[0].font.size = Pt(10)
            
            left_date = left_cell.add_paragraph(f'Дата: {pab["doc_date"]}')
            left_date.runs[0].font.size = Pt(9)
            
            # Правая колонка - Ответственный за выполнение
            right_cell = sig_table.rows[0].cells[1]
            
            right_label = right_cell.paragraphs[0]
            right_label.add_run('Ответственный за выполнение:').bold = True
            right_label.runs[0].font.size = Pt(10)
            
            responsible_person = pab['observations'][0]['responsible_person'] if pab['observations'] else '—'
            right_name = right_cell.add_paragraph(responsible_person)
            right_name.runs[0].font.size = Pt(10)
            
            right_cell.add_paragraph()
            right_cell.add_paragraph()
            
            right_sign = right_cell.add_paragraph('Подпись ______________')
            right_sign.runs[0].font.size = Pt(10)
            
            right_date = right_cell.add_paragraph('Дата: __________')
            right_date.runs[0].font.size = Pt(9)
            
            doc.add_page_break()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        docx_base64 = base64.b64encode(buffer.read()).decode('utf-8')
        
        filename = f"PAB_{pabs_with_obs[0]['doc_number']}" if len(pabs_with_obs) == 1 else "PAB_Multiple"
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'Access-Control-Allow-Origin': '*',
                'Content-Disposition': f'attachment; filename="{filename}.docx"'
            },
            'body': docx_base64,
            'isBase64Encoded': True
        }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': str(e)}),
            'isBase64Encoded': False
        }