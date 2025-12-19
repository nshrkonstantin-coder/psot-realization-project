import json
import os
from typing import Dict, Any
from datetime import datetime
import psycopg2
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import boto3

def send_notification(cur, user_id: int, title: str, message: str, notification_type: str = 'info'):
    '''Отправка уведомления пользователю и администраторам в локальный чат'''
    # Получаем данные пользователя
    cur.execute(f"""
        SELECT fio, position, organization_id 
        FROM t_p80499285_psot_realization_pro.users 
        WHERE id = {user_id}
    """)
    user_data = cur.fetchone()
    
    if not user_data:
        return
    
    user_fio, user_position, org_id = user_data
    user_fio_esc = str(user_fio).replace("'", "''") if user_fio else ''
    user_position_esc = str(user_position).replace("'", "''") if user_position else ''
    title_esc = str(title).replace("'", "''")
    message_esc = str(message).replace("'", "''")
    
    # Используем стандартное название организации
    org_name = 'АО "ГРК "Западная"'
    
    # Отправляем системное уведомление администраторам
    cur.execute(f"""
        INSERT INTO t_p80499285_psot_realization_pro.system_notifications
        (notification_type, severity, title, message, user_id, user_fio, user_position,
         organization_id, organization_name, is_read, created_at)
        VALUES ('{notification_type}', 'info', '{title_esc}', '{message_esc}', 
                {user_id}, '{user_fio_esc}', '{user_position_esc}',
                {org_id if org_id else 'NULL'}, '{org_name}', false, NOW())
    """)
    
    print(f'[Notification] Sent to user {user_id} and admins')

def create_word_document(pab_data: Dict) -> BytesIO:
    '''Создание Word документа ПАБ'''
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Регистрация ПАБ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Основная информация
    doc.add_paragraph(f"Название листа: {pab_data['doc_number']}")
    doc.add_paragraph(f"Номер документа: {pab_data['doc_number']}")
    doc.add_paragraph(f"Дата: {pab_data['doc_date']}")
    doc.add_paragraph(f"ФИО проверяющего: {pab_data['inspector_fio']}")
    doc.add_paragraph(f"Должность проверяющего: {pab_data['inspector_position']}")
    doc.add_paragraph(f"Участок: {pab_data.get('location', '')}")
    doc.add_paragraph(f"Проверяемый объект: {pab_data.get('checked_object', '')}")
    doc.add_paragraph(f"Подразделение: {pab_data.get('department', '')}")
    
    doc.add_paragraph()
    
    # Наблюдения
    for idx, obs in enumerate(pab_data['observations'], 1):
        doc.add_heading(f'Наблюдение №{idx}', level=2)
        doc.add_paragraph(f"Описание: {obs['description']}")
        doc.add_paragraph(f"Категория: {obs.get('category', '')}")
        doc.add_paragraph(f"Вид условий и действий: {obs.get('conditions_actions', '')}")
        doc.add_paragraph(f"Опасные факторы: {obs.get('hazard_factors', '')}")
        doc.add_paragraph(f"Мероприятия: {obs['measures']}")
        doc.add_paragraph(f"Ответственный: {obs.get('responsible_person', '')}")
        doc.add_paragraph(f"Срок: {obs.get('deadline', '')}")
        doc.add_paragraph()
    
    # Сохранение в BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Сохранение ПАБ, создание Word документа и отправка email
    '''
    
    if event.get('httpMethod') == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type, X-User-Id',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    body = json.loads(event.get('body', '{}'))
    
    dsn = os.environ.get('DATABASE_URL')
    conn = psycopg2.connect(dsn)
    cur = conn.cursor()
    schema = 't_p80499285_psot_realization_pro'
    
    # Генерируем номер ПАБ внутри транзакции
    current_year = datetime.now().year
    year_short = str(current_year)[2:]
    
    cur.execute("SELECT counter FROM pab_counter WHERE year = %s FOR UPDATE", (current_year,))
    result = cur.fetchone()
    
    if result:
        counter = result[0] + 1
        cur.execute("UPDATE pab_counter SET counter = %s WHERE year = %s", (counter, current_year))
    else:
        counter = 1
        cur.execute("INSERT INTO pab_counter (year, counter) VALUES (%s, %s)", (current_year, counter))
    
    doc_number = f"ПАБ-{counter}-{year_short}"
    
    headers = event.get('headers', {})
    user_id = headers.get('X-User-Id') or headers.get('x-user-id')
    
    organization_id = None
    if user_id:
        cur.execute("SELECT organization_id FROM t_p80499285_psot_realization_pro.users WHERE id = %s", (int(user_id),))
        org_result = cur.fetchone()
        if org_result:
            organization_id = org_result[0]
    
    if not organization_id:
        organization_id = 1
    
    cur.execute(
        f"""INSERT INTO {schema}.pab_records 
        (doc_number, doc_date, inspector_fio, inspector_position, location, checked_object, department, organization_id, created_at) 
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW()) 
        RETURNING id""",
        (
            doc_number,
            body['date'],
            body['inspectorName'],
            body['inspectorPosition'],
            body.get('area'),
            body.get('inspectedObject'),
            body.get('subdivision'),
            organization_id
        )
    )
    
    pab_id = cur.fetchone()[0]
    
    observations_data = body.get('observations', [])
    
    s3_client = boto3.client('s3',
        endpoint_url='https://bucket.poehali.dev',
        aws_access_key_id=os.environ['AWS_ACCESS_KEY_ID'],
        aws_secret_access_key=os.environ['AWS_SECRET_ACCESS_KEY']
    )
    
    # Собираем email-адреса ответственных лиц
    responsible_emails = set()
    
    for idx, obs in enumerate(observations_data, 1):
        photo_url = None
        
        if obs.get('photo'):
            try:
                import base64
                
                photo_base64 = obs['photo'].split(',')[1] if ',' in obs['photo'] else obs['photo']
                photo_data = base64.b64decode(photo_base64)
                
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                photo_key = f'pab/photos/{doc_number}_obs{idx}_{timestamp}.jpg'
                
                s3_client.put_object(
                    Bucket='files',
                    Key=photo_key,
                    Body=photo_data,
                    ContentType='image/jpeg',
                    ContentDisposition='inline',
                    CacheControl='public, max-age=31536000'
                )
                
                photo_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{photo_key}"
            except Exception as e:
                print(f"Error uploading photo for observation {idx}: {e}")
        
        # Получаем email и telegram ответственного лица
        responsible_person = obs.get('responsible')
        if responsible_person:
            cur.execute(f"SELECT id, email, telegram_chat_id FROM {schema}.users WHERE fio = %s", (responsible_person,))
            resp_row = cur.fetchone()
            if resp_row:
                responsible_user_id = resp_row[0]
                
                if resp_row[1]:
                    responsible_emails.add(resp_row[1])
                
                # Отправляем уведомление в локальный чат (всегда)
                notification_title = f"📋 Новое наблюдение ПАБ #{doc_number}"
                notification_message = f"""Поступило новое наблюдение ПАБ

Номер: {doc_number}
Дата: {body.get('date', '')}

Описание:
{obs.get('description', '')}

Мероприятия:
{obs.get('measures', '')}

Срок выполнения: {obs.get('deadline', 'не указан')}
Ответственный: {responsible_person}"""
                
                send_notification(cur, responsible_user_id, notification_title, notification_message, 'pab')
        
        # Вставляем наблюдение и получаем ID
        cur.execute(
            f"""INSERT INTO {schema}.pab_observations 
            (pab_record_id, observation_number, description, category, 
            conditions_actions, hazard_factors, measures, responsible_person, deadline, photo_url, created_at) 
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
            RETURNING id""",
            (
                pab_id,
                idx,
                obs.get('description'),
                obs.get('category'),
                obs.get('conditions'),
                obs.get('hazards'),
                obs.get('measures'),
                obs.get('responsible'),
                obs.get('deadline'),
                photo_url
            )
        )
        observation_id = cur.fetchone()[0]
        
        # Отправляем Telegram уведомление (если привязан)
        if responsible_person:
            cur.execute(f"SELECT id, telegram_chat_id FROM {schema}.users WHERE fio = %s", (responsible_person,))
            tg_check = cur.fetchone()
            
            if tg_check and tg_check[1]:
                import urllib.request
                import urllib.parse
                
                bot_token = os.environ.get('TELEGRAM_BOT_TOKEN')
                if bot_token:
                    telegram_message = f"""🔔 <b>Новое наблюдение ПАБ</b>

📋 Номер: {doc_number}
📅 Дата: {body.get('date', '')}

⚠️ Описание:
{obs.get('description', '')}

💡 Мероприятия:
{obs.get('measures', '')}

⏰ Срок: {obs.get('deadline', 'не указан')}
👤 Ответственный: {responsible_person}"""
                    
                    send_url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
                    
                    # Добавляем inline-кнопку "Принято"
                    keyboard = {
                        'inline_keyboard': [[
                            {'text': '✅ Принято', 'callback_data': f'accept_pab_{observation_id}'}
                        ]]
                    }
                    
                    payload = json.dumps({
                        'chat_id': tg_check[1],
                        'text': telegram_message,
                        'parse_mode': 'HTML',
                        'reply_markup': keyboard
                    }).encode()
                    
                    try:
                        req = urllib.request.Request(
                            send_url,
                            data=payload,
                            headers={'Content-Type': 'application/json'}
                        )
                        urllib.request.urlopen(req, timeout=5)
                        print(f'[Telegram] Sent PAB notification to user {tg_check[0]}')
                    except Exception as e:
                        print(f'[Telegram] Failed to send PAB: {e}')
    
    conn.commit()
    
    word_doc = create_word_document({
        'doc_number': doc_number,
        'doc_date': body['date'],
        'inspector_fio': body['inspectorName'],
        'inspector_position': body['inspectorPosition'],
        'location': body.get('area'),
        'checked_object': body.get('inspectedObject'),
        'department': body.get('subdivision'),
        'observations': observations_data
    })
    
    doc_filename = f"{doc_number}.docx"
    
    file_key = f'pab/{doc_filename}'
    s3_client.put_object(
        Bucket='files',
        Key=file_key,
        Body=word_doc.getvalue(),
        ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        ContentDisposition=f'inline; filename="{doc_filename}"',
        CacheControl='public, max-age=31536000'
    )
    
    file_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{file_key}"
    
    cur.execute(f"UPDATE {schema}.pab_records SET word_file_url = %s WHERE id = %s", (file_url, pab_id))
    
    user_id = body.get('userId')
    if user_id:
        cur.execute(f"SELECT id FROM {schema}.storage_folders WHERE user_id = %s AND folder_name = 'ПАБ' LIMIT 1", (user_id,))
        folder_row = cur.fetchone()
        
        if not folder_row:
            cur.execute(f"INSERT INTO {schema}.storage_folders (user_id, folder_name, created_at) VALUES (%s, 'ПАБ', NOW()) RETURNING id", (user_id,))
            folder_id = cur.fetchone()[0]
        else:
            folder_id = folder_row[0]
        
        file_size = len(word_doc.getvalue())
        cur.execute(
            f"""INSERT INTO {schema}.storage_files 
            (folder_id, file_name, file_url, file_size, file_type, uploaded_at) 
            VALUES (%s, %s, %s, %s, %s, NOW())""",
            (folder_id, doc_filename, file_url, file_size, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        )
    
    conn.commit()
    
    # Отправка email администратору и пользователю
    admin_email_sent = False
    user_email_sent = False
    email_error = None
    
    smtp_host = os.environ.get('SMTP_HOST', 'smtp.yandex.ru')
    smtp_port = os.environ.get('SMTP_PORT', '587')
    smtp_user = os.environ.get('SMTP_USER', 'ACYBT@yandex.ru')
    smtp_password = os.environ.get('SMTP_PASSWORD')
    admin_email = os.environ.get('ADMIN_EMAIL', 'ACYBT@yandex.ru')
    
    # Получаем email пользователя из базы
    user_email = None
    if user_id:
        cur.execute(f"SELECT email FROM {schema}.users WHERE id = %s", (user_id,))
        user_row = cur.fetchone()
        if user_row and user_row[0]:
            user_email = user_row[0]
    
    if smtp_host and smtp_user and smtp_password:
        pab_url = f"https://lk.psot-realization.pro/pab-view/{pab_id}"
        
        email_body = f"""Зарегистрирован новый ПАБ

Номер: {doc_number}
Дата: {body['date']}
Проверяющий: {body['inspectorName']}
Подразделение: {body.get('subdivision', '')}

Просмотреть ПАБ: {pab_url}

Во вложении находится полный документ.
"""
        
        # Отправка администратору
        if admin_email:
            try:
                msg = MIMEMultipart()
                msg['From'] = smtp_user
                msg['To'] = admin_email
                msg['Subject'] = f"Новая регистрация ПАБ: {doc_number}"
                
                msg.attach(MIMEText(email_body, 'plain', 'utf-8'))
                
                part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                part.set_payload(word_doc.getvalue())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename={doc_filename}')
                msg.attach(part)
                
                server = smtplib.SMTP(smtp_host, int(smtp_port))
                server.set_debuglevel(1)
                server.starttls()
                server.login(smtp_user, smtp_password)
                print(f"✅ SMTP login successful!")
                server.send_message(msg)
                server.quit()
                
                admin_email_sent = True
                print(f"Email sent to admin: {admin_email}")
                
            except Exception as e:
                email_error = f"Admin email error: {str(e)}"
                print(email_error)
        
        # Отправка пользователю
        if user_email and user_email != admin_email:
            try:
                msg_user = MIMEMultipart()
                msg_user['From'] = smtp_user
                msg_user['To'] = user_email
                msg_user['Subject'] = f"Ваш ПАБ зарегистрирован: {doc_number}"
                
                user_email_body = f"""Ваш ПАБ успешно зарегистрирован в системе

Номер: {doc_number}
Дата: {body['date']}
Проверяющий: {body['inspectorName']}
Подразделение: {body.get('subdivision', '')}

Просмотреть ПАБ: {pab_url}

Во вложении находится полный документ.
"""
                
                msg_user.attach(MIMEText(user_email_body, 'plain', 'utf-8'))
                
                part_user = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                part_user.set_payload(word_doc.getvalue())
                encoders.encode_base64(part_user)
                part_user.add_header('Content-Disposition', f'attachment; filename={doc_filename}')
                msg_user.attach(part_user)
                
                server_user = smtplib.SMTP(smtp_host, int(smtp_port))
                server_user.set_debuglevel(1)
                server_user.starttls()
                server_user.login(smtp_user, smtp_password)
                print(f"✅ User SMTP login successful!")
                server_user.send_message(msg_user)
                server_user.quit()
                
                user_email_sent = True
                print(f"Email sent to user: {user_email}")
                
            except Exception as e:
                if email_error:
                    email_error += f"; User email error: {str(e)}"
                else:
                    email_error = f"User email error: {str(e)}"
                print(f"User email error: {e}")
        
        # Отправка ответственным лицам
        for responsible_email in responsible_emails:
            if responsible_email and responsible_email not in [admin_email, user_email]:
                try:
                    msg_resp = MIMEMultipart()
                    msg_resp['From'] = smtp_user
                    msg_resp['To'] = responsible_email
                    msg_resp['Subject'] = f"Назначено мероприятие по ПАБ: {doc_number}"
                    
                    resp_email_body = f"""Вам назначено мероприятие по ПАБ

Номер ПАБ: {doc_number}
Дата: {body['date']}
Проверяющий: {body['inspectorName']}
Подразделение: {body.get('subdivision', '')}

Просмотреть ПАБ и мероприятия: {pab_url}

Во вложении находится полный документ с вашими задачами.
"""
                    
                    msg_resp.attach(MIMEText(resp_email_body, 'plain', 'utf-8'))
                    
                    part_resp = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                    part_resp.set_payload(word_doc.getvalue())
                    encoders.encode_base64(part_resp)
                    part_resp.add_header('Content-Disposition', f'attachment; filename={doc_filename}')
                    msg_resp.attach(part_resp)
                    
                    server_resp = smtplib.SMTP(smtp_host, int(smtp_port))
                    server_resp.starttls()
                    server_resp.login(smtp_user, smtp_password)
                    server_resp.send_message(msg_resp)
                    server_resp.quit()
                    
                    print(f"Email sent to responsible: {responsible_email}")
                    
                except Exception as e:
                    print(f"Error sending email to responsible {responsible_email}: {e}")
    else:
        print("Email credentials not configured")
    
    cur.close()
    conn.close()
    
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*'
        },
        'body': json.dumps({
            'success': True,
            'pabId': pab_id,
            'pabNumber': doc_number,
            'fileUrl': file_url,
            'adminEmailSent': admin_email_sent,
            'userEmailSent': user_email_sent,
            'userEmail': user_email,
            'emailError': email_error
        }, ensure_ascii=False),
        'isBase64Encoded': False
    }