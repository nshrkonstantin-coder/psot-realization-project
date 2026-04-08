import os
import io
import json
import boto3
import psycopg2
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCHEMA = os.environ.get('MAIN_DB_SCHEMA', 't_p80499285_psot_realization_pro')


def get_conn():
    return psycopg2.connect(os.environ['DATABASE_URL'])


def add_page_numbers(doc):
    """Добавляет нумерацию страниц в правом нижнем углу футера."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.clear()
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)


def set_white_background(doc):
    """Устанавливает белый фон страницы и чёрный текст по умолчанию."""
    style = doc.styles['Normal']
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)


def add_code_block(doc, code_text):
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0, 0, 0)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFFFF')
    pPr.append(shd)


def handler(event: dict, context) -> dict:
    """
    GET  — возвращает статистику файлов из БД.
    POST action=save_files — сохраняет файлы в БД (только .py и .tsx).
    POST action=generate   — генерирует Word из файлов в БД (чёрный текст, белый фон, нумерация страниц).
    """

    if event.get('httpMethod') == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': ''
        }

    headers = {'Access-Control-Allow-Origin': '*', 'Content-Type': 'application/json'}

    # GET — статистика
    if event.get('httpMethod') == 'GET':
        conn = get_conn()
        cur = conn.cursor()
        cur.execute(f"SELECT section, COUNT(*) FROM {SCHEMA}.source_files GROUP BY section")
        rows = cur.fetchall()
        cur.close()
        conn.close()
        stats = {row[0]: row[1] for row in rows}
        return {
            'statusCode': 200,
            'headers': headers,
            'body': json.dumps({
                'frontend_count': stats.get('frontend', 0),
                'backend_count': stats.get('backend', 0),
                'total': sum(stats.values())
            })
        }

    body = json.loads(event.get('body') or '{}')
    action = body.get('action', 'generate')

    # POST action=clear — удалить все файлы из БД
    if action == 'clear':
        conn = get_conn()
        cur = conn.cursor()
        cur.execute(f"DELETE FROM {SCHEMA}.source_files")
        conn.commit()
        cur.close()
        conn.close()
        return {
            'statusCode': 200,
            'headers': headers,
            'body': json.dumps({'success': True})
        }

    # POST action=save_files — сохранить файлы в БД (только .py и .tsx)
    if action == 'save_files':
        files = body.get('files', [])
        if not files:
            return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'Нет файлов'})}

        conn = get_conn()
        cur = conn.cursor()
        saved = 0
        for f in files:
            raw_path = f.get('path', '')
            ext = '.' + raw_path.rsplit('.', 1)[-1].lower() if '.' in raw_path else ''
            section = f.get('section', 'backend')
            allowed = ['.py'] if section == 'backend' else ['.tsx']
            if ext not in allowed:
                continue

            path = raw_path
            content = f.get('content', '')
            cur.execute(
                f"INSERT INTO {SCHEMA}.source_files (file_path, file_content, section) "
                f"VALUES (%s, %s, %s) "
                f"ON CONFLICT (file_path) DO UPDATE SET file_content = EXCLUDED.file_content, updated_at = NOW()",
                (path, content, section)
            )
            saved += 1
        conn.commit()
        cur.close()
        conn.close()
        return {
            'statusCode': 200,
            'headers': headers,
            'body': json.dumps({'success': True, 'saved': saved})
        }

    # POST action=generate — генерация Word
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(f"SELECT file_path, file_content, section FROM {SCHEMA}.source_files ORDER BY section, file_path")
    rows = cur.fetchall()
    cur.close()
    conn.close()

    frontend_files = [{'path': r[0], 'content': r[1]} for r in rows if r[2] == 'frontend' and r[0].endswith('.tsx')]
    backend_files  = [{'path': r[0], 'content': r[1]} for r in rows if r[2] == 'backend'  and r[0].endswith('.py')]

    doc = Document()
    set_white_background(doc)
    add_page_numbers(doc)

    title = doc.add_heading('Исходный код программы', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    intro = doc.add_paragraph('Автоматически сгенерированный документ со всеми исходными файлами проекта.')
    for run in intro.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()

    for section_title, section_files in [
        ('Раздел 1. Фронтенд (TypeScript / TSX)', frontend_files),
        ('Раздел 2. Бэкенд (Python)', backend_files),
    ]:
        h = doc.add_heading(section_title, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        if not section_files:
            doc.add_paragraph('Файлы не загружены.')
            continue
        for f in section_files:
            h2 = doc.add_heading(f['path'], level=2)
            for run in h2.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            add_code_block(doc, f['content'] or '// Файл пустой')
            sep = doc.add_paragraph('─' * 100)
            if sep.runs:
                sep.runs[0].font.size = Pt(6)
                sep.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    total = len(frontend_files) + len(backend_files)
    h_total = doc.add_heading('Итого', level=1)
    for run in h_total.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph(f'Фронтенд (.tsx): {len(frontend_files)} файлов')
    doc.add_paragraph(f'Бэкенд (.py): {len(backend_files)} файлов')
    doc.add_paragraph(f'Всего: {total}')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    s3 = boto3.client(
        's3',
        endpoint_url='https://bucket.poehali.dev',
        aws_access_key_id=os.environ['AWS_ACCESS_KEY_ID'],
        aws_secret_access_key=os.environ['AWS_SECRET_ACCESS_KEY']
    )
    key = 'exports/source_code.docx'
    s3.put_object(Bucket='files', Key=key, Body=buffer.read(),
                  ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    cdn_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{key}"

    return {
        'statusCode': 200,
        'headers': headers,
        'body': json.dumps({
            'success': True,
            'url': cdn_url,
            'total_files': total,
            'frontend_count': len(frontend_files),
            'backend_count': len(backend_files),
        })
    }