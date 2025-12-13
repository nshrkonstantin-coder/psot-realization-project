import json
import os
from typing import Dict, Any
import psycopg2
from io import BytesIO
import base64

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Экспорт записей ПК (производственного контроля) в формат Word (DOCX)
    Возвращает файл .docx для скачивания
    '''
    
    method = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'GET, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    if method != 'GET':
        return {
            'statusCode': 405,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Only GET allowed'}),
            'isBase64Encoded': False
        }
    
    query_params = event.get('queryStringParameters', {})
    ids_str = query_params.get('ids', '')
    
    if not ids_str:
        return {
            'statusCode': 400,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Parameter ids is required'}),
            'isBase64Encoded': False
        }
    
    try:
        pc_ids = [int(id_str.strip()) for id_str in ids_str.split(',')]
    except ValueError:
        return {
            'statusCode': 400,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Invalid ids format'}),
            'isBase64Encoded': False
        }
    
    database_url = os.environ.get('DATABASE_URL')
    if not database_url:
        return {
            'statusCode': 500,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Database not configured'}),
            'isBase64Encoded': False
        }
    
    try:
        conn = psycopg2.connect(database_url)
        cursor = conn.cursor()
        
        ids_placeholder = ','.join(['%s'] * len(pc_ids))
        
        query = f"""
            SELECT 
                r.id,
                r.doc_number,
                r.doc_date,
                r.issuer_name,
                r.issuer_position,
                r.recipient_name,
                r.department,
                '' as checked_object
            FROM t_p80499285_psot_realization_pro.production_control_reports r
            WHERE r.id IN ({ids_placeholder})
            ORDER BY r.doc_date DESC
        """
        
        cursor.execute(query, pc_ids)
        pc_records = cursor.fetchall()
        
        if not pc_records:
            cursor.close()
            conn.close()
            return {
                'statusCode': 404,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'No PC records found'}),
                'isBase64Encoded': False
            }
        
        pcs_with_violations = []
        
        for pc_row in pc_records:
            pc_id = pc_row[0]
            
            viol_query = """
                SELECT 
                    v.item_number,
                    v.description,
                    v.measures,
                    v.deadline,
                    COALESCE(u.fio, '') as responsible_person
                FROM t_p80499285_psot_realization_pro.production_control_violations v
                LEFT JOIN t_p80499285_psot_realization_pro.users u ON v.responsible_user_id = u.id
                WHERE v.report_id = %s
                ORDER BY v.item_number
            """
            
            cursor.execute(viol_query, (pc_id,))
            violations = cursor.fetchall()
            
            # Загружаем фотографии для каждого нарушения
            violations_with_photos = []
            for viol in violations:
                violation_id_query = """
                    SELECT id FROM t_p80499285_psot_realization_pro.production_control_violations
                    WHERE report_id = %s AND item_number = %s
                """
                cursor.execute(violation_id_query, (pc_id, viol[0]))
                violation_id_result = cursor.fetchone()
                
                photos = []
                if violation_id_result:
                    violation_id = violation_id_result[0]
                    photos_query = """
                        SELECT photo_url
                        FROM t_p80499285_psot_realization_pro.production_control_photos
                        WHERE violation_id = %s
                        ORDER BY id
                    """
                    cursor.execute(photos_query, (violation_id,))
                    photo_rows = cursor.fetchall()
                    photos = [{'data': photo[0]} for photo in photo_rows]
                
                violations_with_photos.append({
                    'violation_number': viol[0],
                    'description': viol[1],
                    'measures': viol[2],
                    'deadline': str(viol[3]) if viol[3] else '',
                    'responsible_person': viol[4],
                    'photos': photos
                })
            
            # Загружаем подписи
            signatures_query = """
                SELECT user_name, signature_date
                FROM t_p80499285_psot_realization_pro.production_control_signatures
                WHERE report_id = %s
                ORDER BY id
            """
            cursor.execute(signatures_query, (pc_id,))
            signatures_rows = cursor.fetchall()
            signatures = [{'user_name': sig[0], 'date': str(sig[1]) if sig[1] else ''} for sig in signatures_rows]
            
            pcs_with_violations.append({
                'id': pc_row[0],
                'doc_number': pc_row[1],
                'doc_date': str(pc_row[2]) if pc_row[2] else '',
                'inspector_fio': pc_row[3],
                'inspector_position': pc_row[4],
                'location': pc_row[5],
                'department': pc_row[6],
                'checked_object': pc_row[7],
                'violations': violations_with_photos,
                'signatures': signatures
            })
        
        cursor.close()
        conn.close()
        
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        for pc in pcs_with_violations:
            heading1 = doc.add_heading('ПРОТОКОЛ ПРОИЗВОДСТВЕННОГО КОНТРОЛЯ', level=1)
            heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading1.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            
            heading2 = doc.add_heading(pc['doc_number'], level=2)
            heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading2.runs:
                run.font.size = Pt(14)
                run.font.bold = True
            
            doc.add_paragraph()
            
            info_table = doc.add_table(rows=3, cols=4)
            info_table.style = 'Table Grid'
            
            info_data_left = [
                ('Дата проверки:', pc['doc_date']),
                ('Проверяющий:', pc['inspector_fio']),
                ('Должность:', pc['inspector_position'])
            ]
            
            info_data_right = [
                ('Проверяемый объект:', pc['checked_object'] or '—'),
                ('Подразделение:', pc['department']),
                ('Контролирующий:', pc['location'])
            ]
            
            for idx in range(3):
                row = info_table.rows[idx]
                
                label_left = row.cells[0]
                value_left = row.cells[1]
                label_right = row.cells[2]
                value_right = row.cells[3]
                
                label_para_left = label_left.paragraphs[0]
                label_para_left.text = info_data_left[idx][0]
                label_para_left.runs[0].font.bold = True
                label_para_left.runs[0].font.size = Pt(10)
                
                value_para_left = value_left.paragraphs[0]
                value_para_left.text = str(info_data_left[idx][1]) if info_data_left[idx][1] else '—'
                value_para_left.runs[0].font.size = Pt(10)
                
                label_para_right = label_right.paragraphs[0]
                label_para_right.text = info_data_right[idx][0]
                label_para_right.runs[0].font.bold = True
                label_para_right.runs[0].font.size = Pt(10)
                
                value_para_right = value_right.paragraphs[0]
                value_para_right.text = str(info_data_right[idx][1]) if info_data_right[idx][1] else '—'
                value_para_right.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()
            
            obs_heading = doc.add_heading('НАРУШЕНИЯ', level=2)
            obs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in obs_heading.runs:
                run.font.size = Pt(13)
                run.font.bold = True
            
            for viol in pc['violations']:
                viol_table = doc.add_table(rows=1, cols=1)
                viol_table.style = 'Table Grid'
                viol_cell = viol_table.rows[0].cells[0]
                
                viol_title = viol_cell.add_paragraph()
                viol_title.add_run(f"Нарушение №{viol['violation_number']}").bold = True
                viol_title.runs[0].font.size = Pt(11)
                
                fields = [
                    ('Описание нарушения:', viol['description']),
                    ('Мероприятия по устранению:', viol['measures']),
                    ('Ответственный:', viol.get('responsible_person', '')),
                    ('Срок устранения:', viol.get('deadline', '')),
                ]
                
                for field_label, field_value in fields:
                    field_para = viol_cell.add_paragraph()
                    field_para.add_run(f"{field_label} ").bold = True
                    field_para.runs[0].font.size = Pt(10)
                    
                    value_run = field_para.add_run(str(field_value))
                    value_run.font.size = Pt(10)
                
                # Добавляем фотографии
                if viol.get('photos') and len(viol['photos']) > 0:
                    photos_para = viol_cell.add_paragraph()
                    photos_para.add_run('Фото нарушений:').bold = True
                    photos_para.runs[0].font.size = Pt(10)
                    
                    for photo in viol['photos']:
                        try:
                            import base64
                            from io import BytesIO
                            from docx.shared import Inches
                            
                            # Декодируем base64 фото
                            photo_data = photo['data']
                            if photo_data.startswith('data:image'):
                                photo_data = photo_data.split(',')[1]
                            
                            image_bytes = base64.b64decode(photo_data)
                            image_stream = BytesIO(image_bytes)
                            
                            # Добавляем изображение в документ
                            photo_para = viol_cell.add_paragraph()
                            run = photo_para.add_run()
                            run.add_picture(image_stream, width=Inches(3))
                        except Exception as e:
                            print(f'Error adding photo to Word: {str(e)}')
                
                doc.add_paragraph()
            
            doc.add_paragraph()
            
            signatures_heading = doc.add_heading('ПОДПИСИ', level=2)
            signatures_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in signatures_heading.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            
            # Добавляем подпись проверяющего
            inspector_para = doc.add_paragraph()
            inspector_para.add_run('Проверяющий: ').bold = True
            inspector_para.add_run(f'{pc["inspector_fio"]} ').font.size = Pt(10)
            inspector_para.add_run('Подпись ______________').font.size = Pt(10)
            inspector_para.add_run(f'  Дата: {pc["doc_date"]}').font.size = Pt(9)
            
            doc.add_paragraph()
            
            # Добавляем подписи принявших из основной формы
            for sig in pc.get('signatures', []):
                sig_para = doc.add_paragraph()
                sig_para.add_run('Принял: ').bold = True
                sig_para.add_run(f'{sig["user_name"]} ').font.size = Pt(10)
                sig_para.add_run('Подпись ______________').font.size = Pt(10)
                sig_para.add_run(f'  Дата: {sig["date"]}').font.size = Pt(9)
            
            doc.add_page_break()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        docx_base64 = base64.b64encode(buffer.read()).decode('utf-8')
        
        filename = f"PC_{pcs_with_violations[0]['doc_number']}" if len(pcs_with_violations) == 1 else "PC_Multiple"
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'Access-Control-Allow-Origin': '*',
                'Content-Disposition': f'attachment; filename="{filename}.docx"'
            },
            'body': docx_base64,
            'isBase64Encoded': True
        }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': str(e)}),
            'isBase64Encoded': False
        }